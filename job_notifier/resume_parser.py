"""
resume_parser.py
Reads a resume from PDF or DOCX and returns clean plain text.
Supports: .pdf (text-based), .docx, .doc (via python-docx)

Usage:
    from resume_parser import load_resume
    resume_text = load_resume("path/to/resume.pdf")
"""
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)


def load_resume(path: str) -> str:
    """
    Parse a resume file and return its text content.
    Raises FileNotFoundError or ValueError on failure.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Resume not found at: {path}")

    ext = p.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(p)
    elif ext in (".docx", ".doc"):
        return _parse_docx(p)
    elif ext == ".txt":
        return p.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported resume format: {ext}. Use PDF, DOCX, or TXT.")


# ── PDF ─────────────────────────────────────────────────────────────────────

def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for PDF parsing. Run: pip install pypdf")

    reader = PdfReader(str(path))
    pages_text = []

    for page in reader.pages:
        try:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        except Exception as e:
            logger.warning(f"Could not extract text from a page: {e}")

    if not pages_text:
        raise ValueError(
            "PDF appears to be image-based (scanned) — text extraction returned nothing. "
            "Please provide a text-based PDF or convert to DOCX."
        )

    full_text = "\n".join(pages_text)
    return _clean(full_text)


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Run: pip install python-docx")

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables (skills tables are common in resumes)
    for table in doc.tables:
        for row in table.rows:
            row_text = "  |  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    if not paragraphs:
        raise ValueError("DOCX file appears to be empty or unreadable.")

    return _clean("\n".join(paragraphs))


# ── Helpers ──────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Normalise whitespace while preserving paragraph breaks."""
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        line = line.strip()
        if line:
            cleaned.append(line)
        elif cleaned and cleaned[-1] != "":
            cleaned.append("")   # preserve single blank line between sections

    return "\n".join(cleaned).strip()


def preview(path: str, max_chars: int = 500) -> str:
    """Return a short preview of the parsed resume — useful for debugging."""
    text = load_resume(path)
    return text[:max_chars] + ("..." if len(text) > max_chars else "")